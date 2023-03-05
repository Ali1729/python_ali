from docxtpl import DocxTemplate
import pandas as pd
from docx import Document
# from docx2pdf import convert
# import os
# from PyPDF2 import PdfMerger
# from pathlib import Path
from docxcompose.composer import Composer
from docx import Document as Document_compose
from io import FileIO

def combine_word_documents(files):
    combined_document = Document()
    count, number_of_files = 0, len(files)
    for file in files:
        sub_doc = Document(file)

        # Don't add a page break if you've
        # reached the last file.
        if count < number_of_files - 1:
            sub_doc.add_page_break()

        for paragraph in sub_doc.paragraphs:
            text = paragraph.text
            combined_document.add_paragraph(text)
        count += 1

    combined_document.save('combined_word_documents.docx')

def combine_all_docx(files_list):
    filename_master=Document(FileIO('ANNEXURE_TEMPLATE.docx'))
    number_of_sections=len(files_list)
    master = Document_compose(filename_master)
    composer = Composer(master)
    for i in range(0, number_of_sections):
        doc_temp = Document_compose(files_list[i])
        composer.append(doc_temp)
    composer.save("combined_file.docx")

def print_annexure(excel_path):
    tpl = DocxTemplate('ANNEXURE_TEMPLATE.docx')

    df1 = pd.read_excel(excel_path, "TEMPLATESHEET", skiprows=5,
                        nrows=2000, usecols="A:AD", na_filter=False)
    list_of_docs = []
    for index, row in df1.iterrows():

        context = {"STUDENT_NAME": str(row["Applicant Name"]).strip(),
                   "FATHER_NAME": str(row["Fathername"]).strip(),
                   "VILLAGE_NAME": str(row["Fathername"]).strip(),
                   "MANDAL_NAME": str(row["VILLAGE_NAME"]).strip(),
                   "DISTANCE": str(row["MANDAL_NAME"]).strip(),
                   "JOINING_DATE": str(row["Date Of Joining"]).strip(),
                   "ADMIN_NO": str(row["ADMN"]).strip(),
                   "CASTE": str(row["Religion"]).strip(),
                   "SUBCASTE": str(row["Caste"]).strip(),
                   "INCOME": str(row["INCOME"]).strip(),
                   "ATTENDANCE": str(row["ATTENDANCEPERCENTAGE"]).strip()

                   }
        tpl.render(context)
        tpl.save(str(row["ADMN"]).strip()+".docx")
        result_file = str(row["ADMN"]).strip()+".pdf"
        # print(result_file)
        # convert(str(row["ADMN"]).strip()+".docx",str(row["ADMN"]).strip()+".pdf")
        # os.remove(str(row["ADMN"]).strip()+".docx")
        # list_of_docs.append(str(row["ADMN"]).strip()+".docx")
        list_of_docs.append(str(row["ADMN"]).strip()+".pdf")


    # merger = PdfMerger()
    # for pdf in list_of_docs:
    #     merger.append(pdf)

    # merger.write("ANNEXURE_RESULT.pdf")
    # merger.close()

    # combine_all_docx(list_of_docs)

excel_path = "ADMNREGISTER.xlsx"
print_annexure(excel_path)
